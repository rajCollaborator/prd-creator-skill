"""Convert a prd-creator spec .md to .docx. Usage: py md_to_docx.py <src.md> <dst.docx>

Handles the markdown subset the spec template uses: #/##/### headings, pipe tables,
fenced code blocks, bullet and numbered lists, **bold**, *italic*, `inline code`,
--- horizontal rules (skipped). Requires python-docx.
"""
import re
import sys
from docx import Document
from docx.shared import Pt

if len(sys.argv) != 3:
    sys.exit("usage: py md_to_docx.py <src.md> <dst.docx>")
SRC, DST = sys.argv[1], sys.argv[2]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

TOKEN = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+\*)")

def add_runs(par, text):
    for piece in TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.italic = True
        else:
            par.add_run(piece)

def add_code_block(lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(9)

def add_table(rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    ncols = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j in range(ncols):
            txt = row[j] if j < len(row) else ""
            cell = t.cell(i, j)
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], txt)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

with open(SRC, encoding="utf-8") as f:
    src_lines = f.read().splitlines()

i = 0
while i < len(src_lines):
    ln = src_lines[i]
    if ln.startswith("```"):
        j = i + 1
        block = []
        while j < len(src_lines) and not src_lines[j].startswith("```"):
            block.append(src_lines[j]); j += 1
        add_code_block(block)
        i = j + 1
        continue
    if ln.startswith("|"):
        j = i
        rows = []
        while j < len(src_lines) and src_lines[j].startswith("|"):
            if not re.match(r"^\|[\s\-|:]+\|?$", src_lines[j]):
                rows.append(src_lines[j])
            j += 1
        add_table(rows)
        i = j
        continue
    if ln.startswith("### "):
        doc.add_heading(ln[4:], level=3)
    elif ln.startswith("## "):
        doc.add_heading(ln[3:], level=2)
    elif ln.startswith("# "):
        doc.add_heading(ln[2:], level=1)
    elif ln.strip() == "---":
        pass
    elif ln.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, ln[2:])
    elif re.match(r"^\d+\. ", ln):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\d+\. ", "", ln))
    elif ln.strip():
        p = doc.add_paragraph()
        add_runs(p, ln)
    i += 1

doc.save(DST)
print("saved", DST)
